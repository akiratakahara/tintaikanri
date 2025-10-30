"""
契約書テンプレート記入モジュール
実際のテンプレートファイルのテーブルに値を書き込む
※書式・書体を保持したまま値を更新
"""
import os
from pathlib import Path
from typing import Dict
from docx import Document
from docx.oxml import OxmlElement
import shutil


class ContractTemplateFiller:
    """契約書テンプレート記入クラス"""

    def __init__(self):
        self.template_dir = Path(__file__).parent.parent / "templates" / "contracts"

    def _set_cell_text_preserve_format(self, cell, text: str):
        """
        セルのテキストを書式を保持したまま更新

        Args:
            cell: Wordのテーブルセル
            text: 設定するテキスト
        """
        # セルに段落がある場合
        if len(cell.paragraphs) > 0:
            para = cell.paragraphs[0]

            # 既存のrunがある場合は最初のrunのみ更新
            if len(para.runs) > 0:
                # 最初のrun以外を削除
                for i in range(len(para.runs) - 1, 0, -1):
                    para._element.remove(para.runs[i]._element)

                # 最初のrunのテキストを更新（書式は保持される）
                para.runs[0].text = text
            else:
                # runがない場合は新規作成
                para.add_run(text)
        else:
            # 段落がない場合は新規作成
            para = cell.add_paragraph()
            para.add_run(text)

    def fill_teiki_shakuya_keiyaku(self, data: Dict[str, str], output_path: str) -> str:
        """
        定期借家契約書テンプレートに値を記入

        Args:
            data: 入力データの辞書
            output_path: 出力ファイルパス

        Returns:
            生成されたファイルのパス
        """
        # テンプレートファイルをコピー
        template_path = self.template_dir / "定期借家契約書.docx"

        if not template_path.exists():
            raise FileNotFoundError(f"テンプレートが見つかりません: {template_path}")

        # テンプレートを読み込み
        doc = Document(str(template_path))

        # テーブル1: 物件情報
        self._fill_table1_property(doc, data)

        # テーブル2: 契約期間
        self._fill_table2_period(doc, data)

        # テーブル3: 賃料等
        self._fill_table3_rent(doc, data)

        # テーブル4: 賃借人・緊急連絡先
        self._fill_table4_tenant(doc, data)

        # テーブル5: 賃貸人
        self._fill_table5_landlord(doc, data)

        # テーブル6: 管理者等
        self._fill_table6_manager(doc, data)

        # 保存
        doc.save(output_path)
        return output_path

    def _fill_table1_property(self, doc: Document, data: Dict):
        """テーブル1: 物件情報を記入"""
        if len(doc.tables) < 1:
            return

        table = doc.tables[0]

        # 建物名称 (行1, 列2-)
        if '建物名称' in data and len(table.rows) > 0:
            for col_idx in range(2, min(7, len(table.rows[0].cells))):
                self._set_cell_text_preserve_format(
                    table.rows[0].cells[col_idx],
                    data.get('建物名称', '')
                )

        # 建物所在地(住居表示) (行2, 列2-)
        if '建物所在地_住居表示' in data and len(table.rows) > 1:
            for col_idx in range(2, min(7, len(table.rows[1].cells))):
                self._set_cell_text_preserve_format(
                    table.rows[1].cells[col_idx],
                    data.get('建物所在地_住居表示', '')
                )

        # 建物所在地(登記簿) (行3, 列2-)
        if '建物所在地_登記簿' in data and len(table.rows) > 2:
            for col_idx in range(2, min(7, len(table.rows[2].cells))):
                self._set_cell_text_preserve_format(
                    table.rows[2].cells[col_idx],
                    data.get('建物所在地_登記簿', '')
                )

        # 建物構造 (行4)
        if '建物構造' in data and len(table.rows) > 3:
            # 構造に応じてチェックボックス的な処理
            structure = data.get('建物構造', '')
            for col_idx in range(2, min(7, len(table.rows[3].cells))):
                cell_text = table.rows[3].cells[col_idx].text
                if structure in cell_text:
                    # 該当する構造の前に✓を追加
                    self._set_cell_text_preserve_format(
                        table.rows[3].cells[col_idx],
                        f"☑{structure}"
                    )

        # 建物種類 (行5)
        if '建物種類' in data and len(table.rows) > 4:
            kind = data.get('建物種類', '')
            for col_idx in range(2, min(5, len(table.rows[4].cells))):
                cell_text = table.rows[4].cells[col_idx].text
                if kind in cell_text:
                    self._set_cell_text_preserve_format(
                        table.rows[4].cells[col_idx],
                        f"☑{kind}"
                    )

        # 新築年月 (行5, 列6-)
        if '新築年月' in data and len(table.rows) > 4:
            self._set_cell_text_preserve_format(table.rows[4].cells[6], data.get('新築年月', ''))

        # 間取り (行6, 列2-)
        if '間取り' in data and len(table.rows) > 5:
            layout = data.get('間取り', '')
            for col_idx in range(2, min(5, len(table.rows[5].cells))):
                cell_text = table.rows[5].cells[col_idx].text
                if layout in cell_text:
                    self._set_cell_text_preserve_format(
                        table.rows[5].cells[col_idx],
                        f"☑{layout}"
                    )

        # 専有面積 (行6, 列6)
        if '専有面積' in data and len(table.rows) > 5:
            self._set_cell_text_preserve_format(table.rows[5].cells[6], data.get('専有面積', '') + ' ㎡')

        # 附属施設 (行7-11)
        facilities = {
            '駐車場': 7,
            'バイク置場': 8,
            '自転車置場': 9,
            '物置': 10,
            '専用庭': 11
        }

        for facility_name, row_idx in facilities.items():
            if facility_name in data and len(table.rows) > row_idx:
                # "有" or "無" をチェック
                value = data.get(facility_name, '無')
                if value == '有' or value == 'あり':
                    self._set_cell_text_preserve_format(
                        table.rows[row_idx].cells[3],
                        '☑有'
                    )
                else:
                    self._set_cell_text_preserve_format(
                        table.rows[row_idx].cells[3],
                        '☑無'
                    )

    def _fill_table2_period(self, doc: Document, data: Dict):
        """テーブル2: 契約期間を記入"""
        if len(doc.tables) < 2:
            return

        table = doc.tables[1]

        # 契約期間 (行1)
        if len(table.rows) > 0:
            period_text = ''
            if '契約開始日' in data:
                period_text = f"{data['契約開始日']}から"
            if '契約終了日' in data:
                period_text += f"{data['契約終了日']}まで"
            if '契約年数' in data:
                period_text += f"({data['契約年数']}年間)"

            if period_text:
                self._set_cell_text_preserve_format(table.rows[0].cells[0], period_text)

        # 鍵引渡し時期 (行2)
        if '鍵引渡し日' in data and len(table.rows) > 1:
            self._set_cell_text_preserve_format(table.rows[1].cells[1], data.get('鍵引渡し日', ''))

    def _fill_table3_rent(self, doc: Document, data: Dict):
        """テーブル3: 賃料等を記入"""
        if len(doc.tables) < 3:
            return

        table = doc.tables[2]

        # 賃料 (行1, 列1)
        if '賃料' in data and len(table.rows) > 0:
            self._set_cell_text_preserve_format(table.rows[0].cells[1], data.get('賃料', '') + '円')

        # 共益費 (行1, 列2)
        if '共益費' in data and len(table.rows) > 0:
            self._set_cell_text_preserve_format(table.rows[0].cells[2], data.get('共益費', '') + '円')

        # 敷金 (行2, 列1)
        if '敷金' in data and len(table.rows) > 1:
            self._set_cell_text_preserve_format(table.rows[1].cells[1], data.get('敷金', '') + '円')

        # 礼金 (行2, 列2)
        if '礼金' in data and len(table.rows) > 1:
            self._set_cell_text_preserve_format(table.rows[1].cells[2], data.get('礼金', '') + '円')

        # 賃料支払方法 - 振込先 (行9)
        if len(table.rows) > 8:
            bank_info = ''
            if '銀行名' in data:
                bank_info += data['銀行名']
            if '支店名' in data:
                bank_info += ' ' + data['支店名']
            if '口座種別' in data:
                bank_info += ' ' + data['口座種別']
            if '口座番号' in data:
                bank_info += ' ' + data['口座番号']
            if '口座名義人' in data:
                bank_info += ' ' + data['口座名義人']

            if bank_info:
                # 振込先情報を適切なセルに記入
                self._set_cell_text_preserve_format(table.rows[8].cells[5], bank_info)

        # 支払日 (行10)
        if '賃料支払日' in data and len(table.rows) > 9:
            self._set_cell_text_preserve_format(table.rows[9].cells[3], f"毎月{data['賃料支払日']}日")

    def _fill_table4_tenant(self, doc: Document, data: Dict):
        """テーブル4: 賃借人・緊急連絡先を記入"""
        if len(doc.tables) < 4:
            return

        table = doc.tables[3]

        # 賃借人氏名 (行1)
        if '賃借人氏名' in data and len(table.rows) > 0:
            for col_idx in range(1, min(7, len(table.rows[0].cells))):
                self._set_cell_text_preserve_format(
                    table.rows[0].cells[col_idx],
                    data.get('賃借人氏名', '')
                )

        # 賃借人との関係 (行2)
        if '賃借人関係' in data and len(table.rows) > 1:
            self._set_cell_text_preserve_format(table.rows[1].cells[1], data.get('賃借人関係', '本人'))

        # メールアドレス (行3)
        if '賃借人メール' in data and len(table.rows) > 2:
            self._set_cell_text_preserve_format(table.rows[2].cells[1], data.get('賃借人メール', ''))

        # 固定電話 (行4)
        if '賃借人電話' in data and len(table.rows) > 3:
            self._set_cell_text_preserve_format(table.rows[3].cells[1], data.get('賃借人電話', ''))

        # 携帯電話 (行5)
        if '賃借人携帯' in data and len(table.rows) > 4:
            self._set_cell_text_preserve_format(table.rows[4].cells[1], data.get('賃借人携帯', ''))

        # FAX (行6)
        if '賃借人FAX' in data and len(table.rows) > 5:
            self._set_cell_text_preserve_format(table.rows[5].cells[1], data.get('賃借人FAX', ''))

        # 緊急連絡先 - 連帯保証人氏名 (行7)
        if '連帯保証人氏名' in data and len(table.rows) > 6:
            self._set_cell_text_preserve_format(table.rows[6].cells[1], data.get('連帯保証人氏名', ''))

        # 連帯保証人生年月日 (行7, 列3-4)
        if '連帯保証人生年月日' in data and len(table.rows) > 6:
            self._set_cell_text_preserve_format(table.rows[6].cells[3], data.get('連帯保証人生年月日', ''))

    def _fill_table5_landlord(self, doc: Document, data: Dict):
        """テーブル5: 賃貸人を記入"""
        if len(doc.tables) < 5:
            return

        table = doc.tables[4]

        # 賃貸人氏名 (行1, 列1)
        if '賃貸人氏名' in data and len(table.rows) > 0:
            self._set_cell_text_preserve_format(table.rows[0].cells[1], data.get('賃貸人氏名', ''))

        # 賃貸人住所 (行2, 列1)
        if '賃貸人住所' in data and len(table.rows) > 1:
            self._set_cell_text_preserve_format(table.rows[1].cells[1], data.get('賃貸人住所', ''))

    def _fill_table6_manager(self, doc: Document, data: Dict):
        """テーブル6: 管理者等を記入"""
        if len(doc.tables) < 6:
            return

        table = doc.tables[5]

        # 管理者名称 (行1, 列1)
        if '管理者名称' in data and len(table.rows) > 0:
            self._set_cell_text_preserve_format(table.rows[0].cells[1], data.get('管理者名称', ''))

        # 管理者所在地 (行2, 列1)
        if '管理者所在地' in data and len(table.rows) > 1:
            self._set_cell_text_preserve_format(table.rows[1].cells[1], data.get('管理者所在地', ''))

        # 管理者電話番号 (行2, 列2)
        if '管理者電話' in data and len(table.rows) > 1:
            self._set_cell_text_preserve_format(table.rows[1].cells[2], data.get('管理者電話', ''))

        # 宅建業登録番号 (行3)
        if '宅建業登録番号' in data and len(table.rows) > 2:
            self._set_cell_text_preserve_format(table.rows[2].cells[1], data.get('宅建業登録番号', ''))

        # 保証協会登録番号 (行4)
        if '保証協会登録番号' in data and len(table.rows) > 3:
            self._set_cell_text_preserve_format(table.rows[3].cells[1], data.get('保証協会登録番号', ''))


# シングルトンインスタンス
_filler = None


def get_contract_filler() -> ContractTemplateFiller:
    """ContractTemplateFillerのインスタンスを取得"""
    global _filler
    if _filler is None:
        _filler = ContractTemplateFiller()
    return _filler
