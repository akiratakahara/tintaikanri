"""
追加書類生成モジュール
- 契約金明細書
- 預かり証
"""

from pathlib import Path
from docx import Document
import openpyxl
from typing import Dict
from datetime import datetime


class AdditionalDocumentFiller:
    """追加書類（契約金明細書、預かり証）のテンプレート記入クラス"""

    def __init__(self, template_dir: str = "templates/contracts"):
        self.template_dir = Path(template_dir)

    def _set_cell_text_preserve_format(self, cell, text: str):
        """セルのテキストを書式を保持したまま更新（Wordテーブル用）"""
        if len(cell.paragraphs) > 0:
            para = cell.paragraphs[0]
            if len(para.runs) > 0:
                for i in range(len(para.runs) - 1, 0, -1):
                    para._element.remove(para.runs[i]._element)
                para.runs[0].text = text
            else:
                para.add_run(text)
        else:
            para = cell.add_paragraph()
            para.add_run(text)

    def _set_paragraph_text_preserve_format(self, paragraph, text: str):
        """段落のテキストを書式を保持したまま更新"""
        if len(paragraph.runs) > 0:
            # 最初のrun以外を削除
            for i in range(len(paragraph.runs) - 1, 0, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            # 最初のrunのテキストを更新（書式は保持される）
            paragraph.runs[0].text = text
        else:
            # runがない場合は新規追加
            paragraph.add_run(text)

    def fill_contract_statement(self, data: Dict[str, str], output_path: str) -> str:
        """
        契約金明細書を生成

        必要なデータ:
        - 日付, 賃借人氏名
        - 物件名, 号室, 所在地
        - 月額賃料, 月額共益費, 敷金, 礼金
        - 契約開始日, 契約終了日, 契約年数
        - 保証会社委託料, 仲介手数料
        - 銀行名, 支店名, 口座番号, 口座名義人
        """
        template_path = self.template_dir / "契約金明細書.docx"
        doc = Document(str(template_path))

        # 日付と賃借人氏名（段落）
        if len(doc.paragraphs) >= 3:
            # 日付（段落1）
            if '日付' in data:
                self._set_paragraph_text_preserve_format(doc.paragraphs[1], data['日付'])

            # 賃借人氏名（段落2）
            if '賃借人氏名' in data:
                self._set_paragraph_text_preserve_format(doc.paragraphs[2], f"{data['賃借人氏名']}様")

        # テーブル1: 物件情報
        if len(doc.tables) > 0:
            table = doc.tables[0]
            if len(table.rows) >= 2:
                # 物件名と号室
                if '物件名' in data:
                    self._set_cell_text_preserve_format(table.rows[0].cells[1], data['物件名'])
                if '号室' in data:
                    self._set_cell_text_preserve_format(table.rows[0].cells[3], data['号室'])

                # 所在地（3セルに同じ値を設定）
                if '所在地' in data:
                    for col_idx in [1, 2, 3]:
                        self._set_cell_text_preserve_format(table.rows[1].cells[col_idx], data['所在地'])

        # テーブル2: 賃貸条件
        if len(doc.tables) > 1:
            table = doc.tables[1]
            if len(table.rows) >= 3:
                # 月額賃料
                if '月額賃料' in data:
                    self._set_cell_text_preserve_format(table.rows[0].cells[1], data['月額賃料'])
                # 敷金
                if '敷金' in data:
                    self._set_cell_text_preserve_format(table.rows[0].cells[3], data['敷金'])

                # 月額共益費
                if '月額共益費' in data:
                    self._set_cell_text_preserve_format(table.rows[1].cells[1], data['月額共益費'])
                # 礼金
                if '礼金' in data:
                    self._set_cell_text_preserve_format(table.rows[1].cells[3], data['礼金'])

                # 契約期間
                if '契約開始日' in data and '契約終了日' in data:
                    period_text = f"{data['契約開始日']}〜{data['契約終了日']}"
                    self._set_cell_text_preserve_format(table.rows[2].cells[1], period_text)
                    self._set_cell_text_preserve_format(table.rows[2].cells[2], period_text)

                # 契約年数
                if '契約年数' in data:
                    self._set_cell_text_preserve_format(table.rows[2].cells[3], f"{data['契約年数']}年間")

        # テーブル3: 契約金明細
        if len(doc.tables) > 2:
            table = doc.tables[2]

            # 計算用の値を取得
            try:
                rent = int(data.get('賃料', '0').replace(',', '').replace('円', ''))
                common_fee = int(data.get('共益費', '0').replace(',', '').replace('円', ''))
                deposit = int(data.get('敷金額', '0').replace(',', '').replace('円', ''))
                key_money = int(data.get('礼金額', '0').replace(',', '').replace('円', ''))
                guarantee_fee = int(data.get('保証会社委託料', '0').replace(',', '').replace('円', ''))
                brokerage_fee = int(data.get('仲介手数料', '0').replace(',', '').replace('円', ''))
                brokerage_tax = int(brokerage_fee * 0.1)  # 内税10%

                # 各行に金額を設定
                if len(table.rows) >= 9:
                    # 敷金
                    self._set_cell_text_preserve_format(table.rows[0].cells[2], f"{deposit:,}円")
                    # 礼金
                    self._set_cell_text_preserve_format(table.rows[1].cells[2], f"{key_money:,}円")
                    # 賃料
                    month = data.get('初月', '９月')
                    self._set_cell_text_preserve_format(table.rows[2].cells[0], f"{month}分賃料")
                    self._set_cell_text_preserve_format(table.rows[2].cells[2], f"{rent:,}円")
                    # 共益費
                    self._set_cell_text_preserve_format(table.rows[3].cells[0], f"{month}分共益費")
                    self._set_cell_text_preserve_format(table.rows[3].cells[2], f"{common_fee:,}円")
                    # 保証会社委託料
                    self._set_cell_text_preserve_format(table.rows[4].cells[2], f"{guarantee_fee:,}円")

                    # 小計
                    subtotal = deposit + key_money + rent + common_fee + guarantee_fee
                    self._set_cell_text_preserve_format(table.rows[5].cells[2], f"{subtotal:,}円")

                    # 仲介手数料
                    self._set_cell_text_preserve_format(table.rows[6].cells[0], f"仲介手数料")
                    self._set_cell_text_preserve_format(table.rows[6].cells[1], f"内税{brokerage_tax:,}円")
                    self._set_cell_text_preserve_format(table.rows[6].cells[2], f"{brokerage_fee:,}円")

                    # 合計
                    total = subtotal + brokerage_fee
                    self._set_cell_text_preserve_format(table.rows[8].cells[2], f"{total:,}円")

            except Exception as e:
                print(f"金額計算エラー: {e}")

        # テーブル4: 振込先情報
        if len(doc.tables) > 3:
            table = doc.tables[3]
            if len(table.rows) >= 6:
                # 銀行名
                if '銀行名' in data:
                    self._set_cell_text_preserve_format(table.rows[1].cells[1], data['銀行名'])
                # 支店名
                if '支店名' in data:
                    self._set_cell_text_preserve_format(table.rows[2].cells[1], data['支店名'])
                # 口座番号
                if '口座番号' in data:
                    self._set_cell_text_preserve_format(table.rows[4].cells[1], data['口座番号'])
                # 口座名義人
                if '口座名義人' in data:
                    self._set_cell_text_preserve_format(table.rows[5].cells[1], data['口座名義人'])

        doc.save(output_path)
        return output_path

    def fill_deposit_receipt(self, data: Dict[str, str], output_path: str) -> str:
        """
        預かり証（Excel）を生成

        必要なデータ:
        - 賃借人氏名
        - 物件住所
        - 鍵番号リスト（配列）
        - 日付
        """
        template_path = self.template_dir / "預り証.xlsx"
        wb = openpyxl.load_workbook(str(template_path))
        ws = wb.active

        # 賃借人氏名（D10）
        if '賃借人氏名' in data:
            ws['D10'] = data['賃借人氏名']

        # 物件住所と鍵の説明（D17）
        if '物件住所' in data:
            ws['D17'] = f"{data['物件住所']}　鍵\n下記のとおり確かにお預りしました。"

        # 鍵番号リスト（D23, D25, D27, D29...）
        if '鍵リスト' in data and isinstance(data['鍵リスト'], list):
            key_rows = [23, 25, 27, 29]  # 鍵番号を入れる行
            for idx, key_info in enumerate(data['鍵リスト'][:4]):  # 最大4つ
                if idx < len(key_rows):
                    row = key_rows[idx]
                    # 鍵番号
                    ws[f'D{row}'] = key_info.get('番号', '')
                    # 数量
                    ws[f'I{row}'] = key_info.get('数量', 1)

        # 日付（D38）
        if '日付' in data:
            ws['D38'] = data['日付']
        else:
            # デフォルトで今日の日付
            today = datetime.now()
            ws['D38'] = f"令和　 　　{today.year - 2018}年　　　{today.month}月　 　　{today.day}日"

        wb.save(output_path)
        return output_path


def get_additional_document_filler():
    """AdditionalDocumentFillerのインスタンスを取得"""
    return AdditionalDocumentFiller()
